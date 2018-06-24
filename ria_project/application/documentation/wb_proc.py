from django.conf import settings
from application import models
import datetime
import sys
import os

# !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
# from mailmerge import MailMerge

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE

from io import StringIO

FILE = os.path.join(
    settings.BASE_DIR,
    'application', 'documentation', 'reports', 'temp_doc.docx')

PORTRAIT = os.path.join(
    settings.BASE_DIR,
    'application', 'documentation', 'reports', 'template_p.docx')

LANDSCAPE = os.path.join(
    settings.BASE_DIR,
    'application', 'documentation', 'reports', 'template_l.docx')

def formatted_date(date):
    parsed_date = datetime.datetime.strptime(str(date), "%Y-%m-%d")
    return parsed_date.strftime("%d.%m.%Y")

def year_diff(f_date):
    delta = datetime.date.today() - f_date
    print(delta)
    return 1 + delta.days // 365

def today():
    date = datetime.date.today()
    return formatted_date(date)

def year():
    return today().split('.')[2]

# def write_cell(cell_code, value, name):
    # SCRIPT_DIR = os.path.dirname(__file__)
    #
    # PROJECT_PATH = os.path.join(SCRIPT_DIR, os.pardir)
    # PROJECT_PATH = os.path.abspath(PROJECT_PATH)
    #
    # STATIC_DIR = os.path.join(PROJECT_PATH, 'static')
    #
    # FILE = os.path.join(settings.STATIC_PATH, name)
    #
    # print(FILE)
    #
    # wb = load_workbook(filename = FILE)
    # sheet_ranges = wb.active
    #
    # wb.active[cell_code] = str(value)
    #
    # wb.save(filename = FILE)

# информация по оплаченным пошлинам (таблица);
def payments_docx(year=year(), supported=True):
    payments = models.Payment.objects.filter(
        payment_date__range=[str(year) + '-01-01', str(year) + '-12-31'])
    grounds = models.Ground.objects.all().order_by('ground_code')
    header_list = ['№', 'Наименование типа РИД', 'Номер охранного документа',
                   'Наименование пошлины', 'Сумма оплаты',
                   'Дата служебной или оплаты', 'Название РИД']

    doc = Document(LANDSCAPE)

    # section = doc.sections[-1]
    # new_width, new_height = section.page_height, section.page_width
    # section.orientation = WD_ORIENT.LANDSCAPE
    # section.page_width = new_width
    # section.page_height = new_height

    text = 'Информация по оплаченным пошлинам за {0} год'.format(year)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font = p.add_run(text).font
    font.bold = True
    font.size = Pt(10)
    font.name = 'Arial'

    doc.add_paragraph('')

    all_rows = 0
    global_sum = 0
    for g in grounds:
        local_sum = 0
        doc.add_paragraph('Площадка №' + str(g.ground_code))

        table_header = doc.add_table(1, len(header_list))

        # table_header.style = WD_STYLE_TYPE.TABLE
        column_number = 0
        for text in header_list:
            cell = table_header.cell(0, column_number).paragraphs[0]
            font = cell.add_run(text).font
            font.bold = True
            column_number += 1

        row_number = 0
        for p in payments:
            ip = p.intellectual_property
            if not ip.is_request and (ip.ground == g):
                if supported:
                    if ip.is_supported:
                        duty = p.duty

                        row_number += 1
                        local_sum += duty.size

                        table_row = doc.add_table(1, 7)

                        cell = table_row.cell(0, 0).paragraphs[0]
                        cell.add_run(str(row_number))
                        cell = table_row.cell(0, 1).paragraphs[0]
                        cell.add_run(str(ip.type_fk.name))
                        cell = table_row.cell(0, 2).paragraphs[0]
                        cell.add_run(str(ip.protection_title))
                        cell = table_row.cell(0, 3).paragraphs[0]
                        cell.add_run(str(duty.name))
                        cell = table_row.cell(0, 4).paragraphs[0]
                        cell.add_run(str(duty.size) + 'руб.')
                        cell = table_row.cell(0, 5).paragraphs[0]
                        cell.add_run(str(formatted_date(p.payment_date)))
                        cell = table_row.cell(0, 6).paragraphs[0]
                        cell.add_run(str(ip.name))
                else:
                    duty = p.duty

                    row_number += 1
                    local_sum += duty.size

                    table_row = doc.add_table(1, 7)

                    cell = table_row.cell(0, 0).paragraphs[0]
                    cell.add_run(str(row_number))
                    cell = table_row.cell(0, 1).paragraphs[0]
                    cell.add_run(str(ip.type_fk.name))
                    cell = table_row.cell(0, 2).paragraphs[0]
                    cell.add_run(str(ip.protection_title))
                    cell = table_row.cell(0, 3).paragraphs[0]
                    cell.add_run(str(duty.name))
                    cell = table_row.cell(0, 4).paragraphs[0]
                    cell.add_run(str(duty.size) + 'руб.')
                    cell = table_row.cell(0, 5).paragraphs[0]
                    cell.add_run(str(formatted_date(p.payment_date)))
                    cell = table_row.cell(0, 6).paragraphs[0]
                    cell.add_run(str(ip.name))

        t = doc.add_table(1, 2)
        c_1 = t.cell(0, 0).paragraphs[0]
        c_1.add_run('Всего оплачено патентов на площадке №{0}: {1}'.format(
                    g.ground_code, row_number))
        c_2 = t.cell(0, 1).paragraphs[0]
        c_2.add_run('Оплачено: {0}'.format(local_sum))

        doc.add_paragraph('')

        global_sum += local_sum
        all_rows += row_number



    t = doc.add_table(1, 2)
    c_1 = t.cell(0, 0).paragraphs[0]
    c_1.add_run('Всего патентов оплачено: {0}'.format(all_rows))
    c_2 = t.cell(0, 1).paragraphs[0]
    c_2.add_run('Итоговая сумма оплаты: {0}'.format(global_sum))

    doc.save(FILE)
    doc = open(FILE,"rb")

    return doc

# TODO: ОТЧЕТЫ W/E
# список патентов по оплате пошлин на поддержания (сортировка по номеру патента);
def maintenance_of_patents_docx(sorting_field=-1):
    ip_list = models.IntellectualProperty.objects.filter(is_supported=True, is_request=False)
    grounds = models.Ground.objects.all()
    payments = models.Payment.objects.all()
    header_list = ['№', 'Наименование типа РИД', 'Номер охранного документа',
                   'Номер заявки', 'Дата приоритета', 'Год поддержки',
                   'Сумма оплаты', 'Дата служебной или оплаты', 'Название РИД',
                   'Авторы', 'Примечание']

    doc = Document(LANDSCAPE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font = p.add_run('Перечень патентов для поддержки').font
    font.bold = True
    font.size = Pt(10)
    font.name = 'Arial'

    doc.add_paragraph('')

    all_rows = 0
    global_sum = 0
    for g in grounds:
        local_sum = 0
        doc.add_paragraph('Площадка №' + str(g.ground_code))

        table_header = doc.add_table(1, len(header_list))
        # table_header.style = doc.styles['TableGrid']
        column_number = 0
        for text in header_list:
            cell = table_header.cell(0, column_number).paragraphs[0]
            cell.add_run(text)
            column_number += 1

        row_number = 0
        for ip in ip_list.filter(ground=g.id).order_by(sorting_field):
            diff = year_diff(ip.grant_date)
            duty = models.Duty.objects.get(
                intellectual_property_type=ip.type_fk, date=diff)
            payment = payments.filter(
                intellectual_property=ip.id, duty=duty.id)
            if len(payment) < 1:
                row_number += 1
                local_sum += duty.size
                creators = ''

                for creator in ip.creators.all():
                    creators += str(creator).split(',')[0] + ', '

                table_row = doc.add_table(1, 11)
                table_row.style = doc.styles['TG']

                cell = table_row.cell(0, 0).paragraphs[0]
                cell.add_run(str(row_number))
                cell = table_row.cell(0, 1).paragraphs[0]
                cell.add_run(str(ip.type_fk.name))
                cell = table_row.cell(0, 2).paragraphs[0]
                cell.add_run(str(ip.protection_title))
                cell = table_row.cell(0, 3).paragraphs[0]
                cell.add_run(str(ip.request_number))
                cell = table_row.cell(0, 4).paragraphs[0]
                cell.add_run(str(ip.priority_date))
                cell = table_row.cell(0, 5).paragraphs[0]
                cell.add_run(str(diff))
                cell = table_row.cell(0, 6).paragraphs[0]
                cell.add_run(str(duty.size) + 'руб.')
                cell = table_row.cell(0, 7).paragraphs[0]
                cell.add_run('')
                cell = table_row.cell(0, 8).paragraphs[0]
                cell.add_run(str(ip.name))
                cell = table_row.cell(0, 9).paragraphs[0]
                cell.add_run(str(creators[0:-2]))
                cell = table_row.cell(0, 10).paragraphs[0]
                cell.add_run('')


        t = doc.add_table(1, 2)
        c_1 = t.cell(0, 0).paragraphs[0]
        c_1.add_run('Всего патентов на площадке №{0}: {1}'.format(
                    g.ground_code, row_number))
        c_2 = t.cell(0, 1).paragraphs[0]
        c_2.add_run('Требуемая сумма: {0}'.format(local_sum))
        global_sum += local_sum
        all_rows += row_number

    doc.add_paragraph('')
    t = doc.add_table(1, 2)
    c_1 = t.cell(0, 0).paragraphs[0]
    c_1.add_run('Всего патентов: {0}'.format(all_rows))
    c_2 = t.cell(0, 1).paragraphs[0]
    c_2.add_run('Всего (требуемая сумма): {0}'.format(global_sum))

    doc.save(FILE)
    doc = open(FILE,"rb")

    return doc

# список действующих патентов;
def actual_patents_docx(ip_type=-1):
    ip_list = models.IntellectualProperty.objects.filter(type_fk=int(ip_type)).filter(is_supported=True)

    grounds = models.Ground.objects.all()
    header_list = ['№', 'Номер охранного документа', 'Номер заявки',
                   'Дата приоритета', 'Название', 'Авторы']

    doc = Document()

    doc.add_paragraph('Реестр поддерживаемых документов')
    doc.add_paragraph("Тип РИД: " + str(ip_type))
    doc.add_paragraph('')

    all_rows = 0
    for g in grounds:
        doc.add_paragraph('Площадка №' + str(g.ground_code))

        table_header = doc.add_table(1, 6)
        column_number = 0
        for text in header_list:
            cell = table_header.cell(0, column_number).paragraphs[0]
            cell.add_run(text)
            column_number += 1

        row_number = 0
        for ip in ip_list.filter(ground=g.id):

            row_number += 1
            creators = ''

            for creator in ip.creators.all():
                creators += str(creator).split(',')[0] + ', '

            table_row = doc.add_table(1, 6)

            cell = table_row.cell(0, 0).paragraphs[0]
            cell.add_run(str(row_number))
            cell = table_row.cell(0, 1).paragraphs[0]
            cell.add_run(str(ip.protection_title))
            cell = table_row.cell(0, 2).paragraphs[0]
            cell.add_run(str(ip.request_number))
            cell = table_row.cell(0, 3).paragraphs[0]
            cell.add_run(str(ip.priority_date))
            cell = table_row.cell(0, 4).paragraphs[0]
            cell.add_run(str(ip.name))
            cell = table_row.cell(0, 5).paragraphs[0]
            cell.add_run(str(creators[0:-2]))


        p = 'Всего документов по площадке №{0} - {1}'.format(g.ground_code, row_number)
        doc.add_paragraph(p)
        all_rows += row_number

    doc.add_paragraph('Итого документов по всем площадкам - ' + str(all_rows))
    doc.add_paragraph('')
    doc.add_paragraph(today())

    doc.save(FILE)
    doc = open(FILE,"rb")

    return doc

# служебная на оплату пошлины (поддержание патента в силе)
def payment_request_docx(doc_number):
    # info = models.PaymentInfo.objects.all()[0]
    # intellectual_property = models.IntellectualProperty.objects.get(
    #     protection_title=doc_number)
    # diff = year_diff(intellectual_property.grant_date)
    # duty = models.Duty.objects.get(
    #     intellectual_property_type=intellectual_property.type_fk,
    #     date=diff)
    # creators = ''
    #
    # for creator in intellectual_property.creators.all():
    #     creators += str(creator).split(',')[0] + ', '
    #
    # paragraph_1 = (
    #     'Прошу дать распоряжение оплатить пошлину в размере {} руб. за {}-й год '
    #     'действия патента на {} №{} по заявке №{} от {}.'.format(
    #     duty.size, diff, intellectual_property.type_fk.name,
    #     intellectual_property.protection_title,
    #     intellectual_property.request_number,
    #     formatted_date(intellectual_property.priority_date)
    # ))
    # paragraph_2 = '"%s" авторов: ' % intellectual_property.name, creators[0:-2]
    #
    #
    # doc = Document()
    #
    # header_table = doc.add_table(1, 2)
    #
    # ht_cell = header_table.cell(0, 0).paragraphs[0].add_run(info.t_1)
    # header_table.cell(0, 0).add_paragraph(today())
    # ht_cell = header_table.cell(0, 1).paragraphs[0]
    # ht_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # ht_cell.add_run(info.t_2)
    #
    # doc.add_paragraph('')
    # doc.add_paragraph('Служебная записка').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #
    # p1 = doc.add_paragraph(paragraph_1)
    # p2 = doc.add_paragraph(paragraph_2)
    #
    # p1.paragraph_format.first_line_indent = Cm(1.25)
    # p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # p2.paragraph_format.first_line_indent = Cm(1.25)
    # p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #
    #
    # doc.add_paragraph('Реквизиты для уплаты пошлин:')
    # table = doc.add_table(10, 2)
    # table.cell(0, 0).paragraphs[0].add_run('Получатель')
    # table.cell(0, 1).paragraphs[0].add_run(info.f_1)
    # table.cell(1, 0).paragraphs[0].add_run('Расчетный счет')
    # table.cell(1, 1).paragraphs[0].add_run(info.f_2)
    # table.cell(2, 0).paragraphs[0].add_run('Банк получателя')
    # table.cell(2, 1).paragraphs[0].add_run(info.f_3)
    # table.cell(3, 0).paragraphs[0].add_run('Корр. счет')
    # table.cell(3, 1).paragraphs[0].add_run(info.f_4)
    # table.cell(4, 0).paragraphs[0].add_run('БИК')
    # table.cell(4, 1).paragraphs[0].add_run(info.f_5)
    # table.cell(5, 0).paragraphs[0].add_run('ИНН')
    # table.cell(5, 1).paragraphs[0].add_run(info.f_6)
    # table.cell(6, 0).paragraphs[0].add_run('КПП')
    # table.cell(6, 1).paragraphs[0].add_run(info.f_7)
    # table.cell(7, 0).paragraphs[0].add_run('ОКАТО')
    # table.cell(7, 1).paragraphs[0].add_run(info.f_8)
    # table.cell(8, 0).paragraphs[0].add_run('ОКПО')
    # table.cell(8, 1).paragraphs[0].add_run(info.f_9)
    # table.cell(9, 0).paragraphs[0].add_run('ОГРН')
    # table.cell(9, 1).paragraphs[0].add_run(info.f_10)
    #
    # doc.add_paragraph('')
    #
    # paragraph_3 = (
    #     '{}, пошлина за {}-й год '
    #     'действия патента на {} №{} по заявке №{} от {}.'.format(
    #         info.i, diff, intellectual_property.type_fk.name,
    #         intellectual_property.protection_title,
    #         intellectual_property.request_number,
    #         formatted_date(intellectual_property.priority_date)
    # ))
    # doc.add_paragraph('В платежном поручении указать:')
    # p3 = doc.add_paragraph(paragraph_3)
    # p3.paragraph_format.first_line_indent = Cm(1.25)
    # p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #
    # footer_table = doc.add_table(1, 2)
    # ht_cell = footer_table.cell(0, 0).paragraphs[0]
    # ht_cell.add_run('Начальник отдела')
    # ht_cell = footer_table.cell(0, 1).paragraphs[0]
    # ht_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # ht_cell.add_run('/{0}/'.format(info.h))

    doc = Document()

    doc.save(FILE)
    doc = open(FILE,"rb")

    return doc

# количество поданных заявок в целом и по площадкам.
def requests_statistics_docx(ip_type=-1, year=year()):
    ip_list = models.IntellectualProperty.objects.filter(
        send_date__range=[str(year) + '-01-01', str(year) + '-12-31'])
    grounds = models.Ground.objects.all()
    header_list = ['№', 'Дата подачи заявки', 'Номер охранного документа', 'Номер заявки',
                   'Дата приоритета', 'Название', 'Авторы', 'Патентообладатели']

    doc = Document()

    doc.add_paragraph('Реестр заявок за {0} год'.format(year))
    doc.add_paragraph('Тип РИД: {0}'.format(str(ip_type)))
    doc.add_paragraph('')

    all_rows = 0
    for g in grounds:
        doc.add_paragraph('Площадка №' + str(g.ground_code))

        table_header = doc.add_table(1, 8)
        column_number = 0
        for text in header_list:
            cell = table_header.cell(0, column_number).paragraphs[0]
            cell.add_run(text)
            column_number += 1

        row_number = 0
        for ip in ip_list.filter(ground=g.id):
            row_number += 1
            creators = ''
            owners = ''

            for creator in ip.creators.all():
                creators += str(creator).split(',')[0] + ', '

            for owner in ip.owners.all():
                owners += str(owner).split(',')[0] + ', '

            table_row = doc.add_table(1, 8)

            cell = table_row.cell(0, 0).paragraphs[0]
            cell.add_run(str(row_number))
            cell = table_row.cell(0, 1).paragraphs[0]
            cell.add_run(str(formatted_date(ip.send_date)))
            cell = table_row.cell(0, 2).paragraphs[0]
            cell.add_run(str(ip.protection_title))
            cell = table_row.cell(0, 3).paragraphs[0]
            cell.add_run(str(ip.request_number))
            cell = table_row.cell(0, 4).paragraphs[0]
            cell.add_run(str(ip.priority_date))
            cell = table_row.cell(0, 5).paragraphs[0]
            cell.add_run(str(ip.name))
            cell = table_row.cell(0, 6).paragraphs[0]
            cell.add_run(str(creators[0:-2]))
            cell = table_row.cell(0, 7).paragraphs[0]
            cell.add_run(str(owners[0:-2]))

        p = 'Всего заявок по площадке №{0} - {1}'.format(g.ground_code, row_number)
        doc.add_paragraph(p)
        all_rows += row_number

    doc.add_paragraph('Итого заявок по всем площадкам - ' + str(all_rows))
    doc.add_paragraph('')
    doc.add_paragraph(today())

    doc.save(FILE)
    doc = open(FILE,"rb")

    return doc

    pass

# список патентов, поданным по грантам
def patents_statistics_docx(contract_type=-1, year=today().split('.')[2]):
    ip_list = models.IntellectualProperty.objects.filter(
              contract_type=int(contract_type)).filter(
              send_date__range=[str(year) + '-01-01', str(year) + '-12-31'])
    contract_type = models.ContractType.objects.get(id=int(contract_type))

    grounds = models.Ground.objects.all()
    header_list = ['№', 'Номер охранного документа', 'Номер заявки',
                   'Дата приоритета', 'Тип РИД', 'Название', 'Авторы']

    doc = Document()

    doc.add_paragraph('Cписок патентов по воду договора: ' + str(contract_type.name))
    doc.add_paragraph('')

    all_rows = 0
    for g in grounds:
        doc.add_paragraph('Площадка №' + str(g.ground_code))

        table_header = doc.add_table(1, 7)
        column_number = 0
        for text in header_list:
            cell = table_header.cell(0, column_number).paragraphs[0]
            cell.add_run(text)
            column_number += 1

        row_number = 0
        for ip in ip_list.filter(ground=g.id):

            row_number += 1
            creators = ''

            for creator in ip.creators.all():
                creators += str(creator).split(',')[0] + ', '

            table_row = doc.add_table(1, 7)

            cell = table_row.cell(0, 0).paragraphs[0]
            cell.add_run(str(row_number))
            cell = table_row.cell(0, 1).paragraphs[0]
            cell.add_run(str(ip.protection_title))
            cell = table_row.cell(0, 2).paragraphs[0]
            cell.add_run(str(ip.request_number))
            cell = table_row.cell(0, 3).paragraphs[0]
            cell.add_run(str(ip.priority_date))
            cell = table_row.cell(0, 4).paragraphs[0]
            cell.add_run(str(contract_type.name))
            cell = table_row.cell(0, 5).paragraphs[0]
            cell.add_run(str(ip.name))
            cell = table_row.cell(0, 6).paragraphs[0]
            cell.add_run(str(creators[0:-2]))


        p = 'Всего документов по площадке №{0} - {1}'.format(g.ground_code, row_number)
        doc.add_paragraph(p)
        all_rows += row_number

    doc.add_paragraph('Итого документов по всем площадкам - ' + str(all_rows))
    doc.add_paragraph('')
    doc.add_paragraph(today())

    doc.save(FILE)
    doc = open(FILE,"rb")

    return doc


def table_23_docx(year):
    ip_list = models.IntellectualProperty.objects.filter(
              priority_date__range=[str(year) + '-01-01', str(year) + '-12-31'],
              is_request=False)

    doc = Document()

    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    text = (
        'ОСНОВНЫЕ ПОКАЗАТЕЛИ РЕЗУЛЬТАТИВНОСТИ ИССЛЕДОВАНИЙ И РАЗРАБОТОК, '
        'КАДРОВОГО ПОТЕНЦИАЛА И ПОДГОТОВКИ КАДРОВ ВЫСШЕЙ КВАЛИФИКАЦИИ ПО '
        'МЕЖДУНАРОДНОЙ СИСТЕМЕ КЛАССИФИКАЦИИ В {} ГОДУ'.format(year)
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font = p.add_run(text).font
    font.bold = True
    font.size = Pt(10)
    font.name = 'Arial'

    header_table = doc.add_table(4, 18, style='Table Grid')

    row = header_table.row_cells(0)
    row.height = Cm(0.5)

    c_2 = header_table.cell(2, 0)
    c_1 = header_table.cell(0, 0).merge(c_2)




    ###########################
    # contract_type = models.ContractType.objects.get(id=int(contract_type))
    #
    # grounds = models.Ground.objects.all()
    # header_list = ['№', 'Номер охранного документа', 'Номер заявки',
    #                'Дата приоритета', 'Тип РИД', 'Название', 'Авторы']
    #
    # doc = Document()
    #
    # doc.add_paragraph('Cписок патентов по воду договора: ' + str(contract_type.name))
    # doc.add_paragraph('')
    #
    # all_rows = 0
    # for g in grounds:
    #     doc.add_paragraph('Площадка №' + str(g.ground_code))
    #
    #     table_header = doc.add_table(1, 7)
    #     column_number = 0
    #     for text in header_list:
    #         cell = table_header.cell(0, column_number).paragraphs[0]
    #         cell.add_run(text)
    #         column_number += 1
    #
    #     row_number = 0
    #     for ip in ip_list.filter(ground=g.id):
    #
    #         row_number += 1
    #         creators = ''
    #
    #         for creator in ip.creators.all():
    #             creators += str(creator).split(',')[0] + ', '
    #
    #         table_row = doc.add_table(1, 7)
    #
    #         cell = table_row.cell(0, 0).paragraphs[0]
    #         cell.add_run(str(row_number))
    #         cell = table_row.cell(0, 1).paragraphs[0]
    #         cell.add_run(str(ip.protection_title))
    #         cell = table_row.cell(0, 2).paragraphs[0]
    #         cell.add_run(str(ip.request_number))
    #         cell = table_row.cell(0, 3).paragraphs[0]
    #         cell.add_run(str(ip.priority_date))
    #         cell = table_row.cell(0, 4).paragraphs[0]
    #         cell.add_run(str(contract_type.name))
    #         cell = table_row.cell(0, 5).paragraphs[0]
    #         cell.add_run(str(ip.name))
    #         cell = table_row.cell(0, 6).paragraphs[0]
    #         cell.add_run(str(creators[0:-2]))
    #
    #
    #     p = 'Всего документов по площадке №{0} - {1}'.format(g.ground_code, row_number)
    #     doc.add_paragraph(p)
    #     all_rows += row_number
    #
    # doc.add_paragraph('Итого документов по всем площадкам - ' + str(all_rows))
    # doc.add_paragraph('')
    # doc.add_paragraph(today())

    doc.save(FILE)
    doc = open(FILE,"rb")

    return doc
